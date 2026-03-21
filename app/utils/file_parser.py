from typing import Optional, List
import os
import io


async def extract_text_from_file(file_content: bytes, file_extension: str) -> str:
    """
    Extract text from various file formats.
    """
    if file_extension == ".pdf":
        return await extract_text_from_pdf(file_content)
    elif file_extension in [".txt", ".md"]:
        return extract_text_from_plain(file_content)
    elif file_extension in [".docx", ".doc"]:
        return await extract_text_from_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


async def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required for PDF processing")
    
    text_parts = []
    
    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    return "\n\n".join(text_parts)


async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX files."""
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required for DOCX processing")
    
    text_parts = []
    
    with io.BytesIO(file_content) as f:
        doc = docx.Document(f)
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
    
    return "\n\n".join(text_parts)


def extract_text_from_plain(file_content: bytes) -> str:
    """Extract text from plain text or markdown files."""
    try:
        return file_content.decode("utf-8")
    except UnicodeDecodeError:
        return file_content.decode("latin-1", errors="ignore")


def validate_file(
    file_name: str, 
    file_size: int, 
    max_size_mb: int = 10,
    allowed_extensions: Optional[List[str]] = None
) -> tuple[bool, Optional[str]]:
    """
    Validate uploaded file.
    
    Returns:
        Tuple of (is_valid, error_message)
    """
    if allowed_extensions is None:
        allowed_extensions = [".pdf", ".txt", ".md", ".docx", ".doc"]
    
    _, ext = os.path.splitext(file_name)
    ext = ext.lower()
    
    if ext not in allowed_extensions:
        return False, f"File type not allowed. Supported types: {', '.join(allowed_extensions)}"
    
    max_size_bytes = max_size_mb * 1024 * 1024
    if file_size > max_size_bytes:
        return False, f"File size exceeds maximum of {max_size_mb}MB"
    
    return True, None


def get_mime_type(file_extension: str) -> str:
    """Get MIME type from file extension."""
    mime_types = {
        ".pdf": "application/pdf",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword"
    }
    return mime_types.get(file_extension.lower(), "application/octet-stream")


def get_file_extension(filename: str) -> str:
    """Extract and return file extension (e.g. '.pdf')."""
    _, ext = os.path.splitext(filename)
    return ext.lower()


def get_file_name(filename: str) -> str:
    """Extract file name without extension."""
    name, _ = os.path.splitext(filename)
    return os.path.basename(name)
